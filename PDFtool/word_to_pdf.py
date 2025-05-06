import os
import shutil
import re
from PySide6.QtWidgets import QFileDialog, QMessageBox
from docx import Document  # Install python-docx library
from reportlab.lib.pagesizes import letter
from reportlab.pdfgen import canvas
from reportlab.lib import colors
from docx.shared import RGBColor

class WordToPDF:
    def __init__(self, ui, parent=None):
        self.ui = ui
        self.parent = parent
        self.word_file_path = None
        self.pdf_file_path = "Word2PDF.pdf"
        self.saved = False

        # Connect buttons
        self.ui.pushButton_28.clicked.connect(self.select_file)
        self.ui.pushButton_29.clicked.connect(self.upload_file)
        self.ui.pushButton_8.clicked.connect(self.save_file)
        self.ui.pushButton_7.clicked.connect(self.discard_file)
        self.ui.pushButton_10.clicked.connect(self.transform_to_pdf)
        self.ui.pushButton_11.clicked.connect(self.download_pdf)
        self.ui.pushButton_33.clicked.connect(self.delete_pdf)

    def select_file(self):
        self.word_file_path, _ = QFileDialog.getOpenFileName(self.parent, "Select Word File", "", "Word Files (*.docx)")
        if self.word_file_path:
            self.ui.lineEditFilePath_5.setText(self.word_file_path)

    def upload_file(self):
        if self.word_file_path:
            self.ui.progressBar_25.setValue(100)  # Simulate upload
            file_name = os.path.basename(self.word_file_path)
            self.ui.textBrowser_4.setText(file_name)

    def save_file(self):
        if self.word_file_path:
            self.saved = True
            self.ui.lcdNumber_4.display(1)

    def discard_file(self):
        self.word_file_path = None
        self.saved = False
        self.ui.lineEditFilePath_5.clear()
        self.ui.textBrowser_4.clear()
        self.ui.lcdNumber_4.display(0)
        self.ui.progressBar_25.setValue(0)
    
    ''' # mit Farbe, aber bad lines break
    def transform_to_pdf(self):
        if self.saved and self.word_file_path:
            self.ui.progressBar_12.setValue(20)  # Simulate processing
            try:
                # Load the Word document
                doc = Document(self.word_file_path)

                # Initialize the PDF canvas
                pdf_canvas = canvas.Canvas(self.pdf_file_path, pagesize=letter)
                left_margin = 50
                top_margin = 750
                line_spacing = 15
                page_width = 550  # Approximate width of the usable area
                y_position = top_margin   # Start from the top of the page
   
                for paragraph in doc.paragraphs:
                    # Iterate over runs to handle formatting
                   for run in paragraph.runs:
                       if y_position < 50:      # Add a new page if y_position exceeds the page limit
                          pdf_canvas.showPage()
                          pdf_canvas.setFont("Times-Roman", 12)
                          y_position = top_margin 
                  
                        # Apply formatting
                       if run.bold and run.italic:
                           pdf_canvas.setFont("Times-BoldItalic", 12)
                       elif run.bold:
                           pdf_canvas.setFont("Times-Bold", 12)
                       elif run.italic:
                           pdf_canvas.setFont("Times-Italic", 12)
                       else:
                           pdf_canvas.setFont("Times-Roman", 12)

                       # Set text color
                       if run.font.color and isinstance(run.font.color.rgb, RGBColor):
                           color = run.font.color.rgb
                           pdf_canvas.setFillColorRGB(color[0] / 255, color[1] / 255, color[2] / 255)
                       else:
                           pdf_canvas.setFillColor(colors.black)

                       # Draw text for each run
                       words = run.text.split(" ")
                       line = ""
                       for word in words:
                           # Check if the current line exceeds the page width
                           test_line = f"{line} {word}".strip()
                           if pdf_canvas.stringWidth(test_line, "Times-Roman", 12) > page_width:
                               # Write the current line and start a new one
                               pdf_canvas.drawString(left_margin, y_position, line.strip())
                               y_position -= line_spacing
                               line = word
                               if y_position < 50:  # Add a new page if y_position exceeds the page limit
                                   pdf_canvas.showPage()
                                   pdf_canvas.setFont("Times-Roman", 12)
                                   y_position = top_margin
                           else:
                               line = test_line

                        # Write the last line of the run
                       if line:
                           pdf_canvas.drawString(left_margin, y_position, line.strip())
                           y_position -= line_spacing

                    # Add spacing for paragraph-level "Enter"
                   y_position -= line_spacing
                   if y_position < 50:
                       pdf_canvas.showPage()
                       pdf_canvas.setFont("Times-Roman", 12)
                       y_position = top_margin

                # Finalize the PDF
                pdf_canvas.save()
                self.ui.progressBar_12.setValue(100)
                self.ui.textBrowser_9.setText("Word2PDF.pdf")
                self.ui.lcdNumber_9.display(1)
                QMessageBox.information(self.parent, "Success", "Word file transformed to PDF successfully!")
            except Exception as e:
                QMessageBox.critical(self.parent, "Error", f"Failed to transform Word to PDF: {str(e)}")
                self.ui.progressBar_12.setValue(0)
                '''
    ''' #schwarz und weiss + lines--gut, aber keine Farbe
    def transform_to_pdf(self):
       if self.saved and self.word_file_path:
           self.ui.progressBar_12.setValue(20)  # Simulate processing
           try:
               # Load the Word document
               doc = Document(self.word_file_path)

               # Initialize the PDF canvas
               pdf_canvas = canvas.Canvas(self.pdf_file_path, pagesize=letter)
               left_margin = 50
               top_margin = 750
               line_spacing = 15
               page_width = 550  # Approximate width of the usable area
               y_position = top_margin

               # Set default font
               pdf_canvas.setFont("Times-Roman", 12)

               for paragraph in doc.paragraphs:
                   # Handle paragraph-level line breaks
                   text = paragraph.text
                   if not text.strip():
                       # Empty paragraph adds a blank line
                       y_position -= line_spacing
                       continue

                    # Wrap text to fit within the page width
                   words = text.split(" ")
                   line = ""
                   for word in words:
                       # Check if the current line exceeds the page width
                       test_line = f"{line} {word}".strip()
                       if pdf_canvas.stringWidth(test_line, "Times-Roman", 12) > page_width:
                           # Write the current line and start a new one
                           pdf_canvas.drawString(left_margin, y_position, line.strip())
                           y_position -= line_spacing
                           line = word
                           if y_position < 50:  # Add a new page if y_position exceeds the page limit
                               pdf_canvas.showPage()
                               pdf_canvas.setFont("Times-Roman", 12)
                               y_position = top_margin
                       else:
                           line = test_line

                   # Write the last line of the paragraph
                   if line:
                       pdf_canvas.drawString(left_margin, y_position, line.strip())
                       y_position -= line_spacing

                   # Check if we need a new page after the paragraph
                   if y_position < 50:
                       pdf_canvas.showPage()
                       pdf_canvas.setFont("Times-Roman", 12)
                       y_position = top_margin
  
                # Finalize the PDF
               pdf_canvas.save()
               self.ui.progressBar_12.setValue(100)
               self.ui.textBrowser_9.setText("Word2PDF.pdf")
               self.ui.lcdNumber_9.display(1)
            
               QMessageBox.information(self.parent, "Success", "Word file transformed to PDF successfully!")
           except Exception as e:
               self.ui.progressBar_12.setValue(0)
               QMessageBox.critical(self.parent, "Error", f"Failed to transform Word to PDF: {str(e)}")
    
               '''
    # final version, still not good, the table can not be transformed
    def transform_to_pdf(self):
       if self.saved and self.word_file_path:
           self.ui.progressBar_12.setValue(20)  # Simulate processing
           try:
               # Load the Word document
               doc = Document(self.word_file_path)

               # Initialize the PDF canvas
               pdf_canvas = canvas.Canvas(self.pdf_file_path, pagesize=letter)
               left_margin = 50
               top_margin = 750
               line_spacing = 15
               page_width = 550  # Approximate width of the usable area
               y_position = top_margin

               # Set default font
               pdf_canvas.setFont("Times-Roman", 12)

               for paragraph in doc.paragraphs:
                    # Handle paragraph-level line breaks
                   if not paragraph.text.strip():
                       y_position -= line_spacing
                       continue

                    # Process each run in the paragraph for formatting
                   x_position = left_margin
                   for run in paragraph.runs:
                       # Apply font style based on formatting
                       if run.bold and run.italic:
                           pdf_canvas.setFont("Times-BoldItalic", 12)
                       elif run.bold:
                           pdf_canvas.setFont("Times-Bold", 12)
                       elif run.italic:
                           pdf_canvas.setFont("Times-Italic", 12)
                       else:
                           pdf_canvas.setFont("Times-Roman", 12)

                        # Set text color if available
                       if run.font.color and isinstance(run.font.color.rgb, RGBColor):
                           color = run.font.color.rgb
                           pdf_canvas.setFillColorRGB(color[0] / 255, color[1] / 255, color[2] / 255)
                       else:
                           pdf_canvas.setFillColor(colors.black)

                       # Get the text content of the run
                       run_text = run.text

                       # Handle line wrapping
                       words = run_text.split(" ")
                       for word in words:
                           test_line = f"{word} "
                           if pdf_canvas.stringWidth(test_line, pdf_canvas._fontname, 12) + x_position > page_width:
                                # Move to the next line if width is exceeded
                               y_position -= line_spacing
                               x_position = left_margin

                           if y_position < 50:  # Add a new page if y_position exceeds the page limit
                               pdf_canvas.showPage()
                               pdf_canvas.setFont("Times-Roman", 12)  # Reset font on new page
                               x_position = left_margin
                               y_position = top_margin

                            # Draw the word and update the x_position
                           pdf_canvas.drawString(x_position, y_position, word)
                           x_position += pdf_canvas.stringWidth(test_line, pdf_canvas._fontname, 12)

                       # Add a space between runs
                       x_position += pdf_canvas.stringWidth(" ", pdf_canvas._fontname, 12)

                   # Move to the next line after the paragraph
                   y_position -= line_spacing

                   # Check if we need a new page after the paragraph
                   if y_position < 50:
                       pdf_canvas.showPage()
                       pdf_canvas.setFont("Times-Roman", 12)
                       y_position = top_margin
 
               # Finalize the PDF
               pdf_canvas.save()
               self.ui.progressBar_12.setValue(100)
               self.ui.textBrowser_9.setText("Word2PDF.pdf")
               self.ui.lcdNumber_9.display(1)

               QMessageBox.information(self.parent, "Success", "Word file transformed to PDF successfully!")
           except Exception as e:
               self.ui.progressBar_12.setValue(0)
               QMessageBox.critical(self.parent, "Error", f"Failed to transform Word to PDF: {str(e)}")

    def download_pdf(self):
        if self.saved and os.path.exists(self.pdf_file_path):
            download_folder = QFileDialog.getExistingDirectory(self.parent, "Select Download Folder")
            if download_folder:
                destination = os.path.join(download_folder, "Word2PDF.pdf")
                try: 
                    shutil.copy(self.pdf_file_path, destination)
                    self.ui.progressBar_13.setValue(100)
                    QMessageBox.information(self.parent, "Success", f"PDF downloaded to {destination}")
                except Exception as e:
                    QMessageBox.critical(self.parent, "Error", f"Failed to download PDF: {str(e)}")
               
    def delete_pdf(self):
        if os.path.exists(self.pdf_file_path):
            try:
               os.remove(self.pdf_file_path)
               self.ui.textBrowser_9.clear()
               self.ui.lcdNumber_9.display(0)
               self.ui.progressBar_12.setValue(0)
               self.ui.progressBar_13.setValue(0)
               QMessageBox.information(self.parent, "Deleted", "PDF file deleted successfully!")
            except Exception as e:
                QMessageBox.critical(self.parent, "Error", f"Failed to delete PDF: {str(e)}")